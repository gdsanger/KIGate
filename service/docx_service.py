"""
DOCX Service for extracting text content from DOCX files
"""
import logging
from typing import List, BinaryIO
from fastapi import UploadFile, HTTPException
import io
from docx import Document

logger = logging.getLogger(__name__)


class DocxService:
    """Service for handling DOCX operations"""
    
    @staticmethod
    async def extract_text_from_docx(docx_file: UploadFile) -> str:
        """
        Extract text content from a DOCX file
        
        Args:
            docx_file: UploadFile object containing DOCX data
            
        Returns:
            str: Extracted text content from all paragraphs
            
        Raises:
            HTTPException: If DOCX cannot be processed
        """
        try:
            # Read the file content
            content = await docx_file.read()
            
            # Create a file-like object from the bytes
            docx_stream = io.BytesIO(content)
            
            # Create Document reader
            document = Document(docx_stream)
            
            # Extract text from all paragraphs
            text_content = []
            paragraph_count = 0
            
            for paragraph in document.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:  # Only add non-empty paragraphs
                    paragraph_count += 1
                    text_content.append(paragraph_text)
            
            # Also extract text from tables if present
            table_count = 0
            for table in document.tables:
                table_count += 1
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append(f"--- Table {table_count} ---")
                    text_content.extend(table_text)
                    text_content.append("")  # Add spacing after table
            
            if not text_content:
                raise HTTPException(
                    status_code=400,
                    detail="No text content could be extracted from the DOCX file"
                )
            
            # Join all content with newlines
            full_text = "\n".join(text_content)
            
            logger.info(f"Successfully extracted text from DOCX with {paragraph_count} paragraphs and {table_count} tables")
            
            return full_text
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to process DOCX file: {str(e)}"
            )
        finally:
            # Reset file pointer for potential reuse
            try:
                if hasattr(docx_file, 'seek'):
                    await docx_file.seek(0)
                elif hasattr(docx_file, 'file') and hasattr(docx_file.file, 'seek'):
                    docx_file.file.seek(0)
            except Exception:
                pass  # Ignore errors during cleanup
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 4000, overlap: int = 200) -> List[str]:
        """
        Split text into chunks with optional overlap for better context preservation
        
        Args:
            text: Input text to chunk
            chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List[str]: List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            
            # Try to break at a sentence or paragraph boundary near the end
            if end < len(text):
                # Look for sentence endings within the last 10% of the chunk
                search_start = max(start, end - chunk_size // 10)
                
                # Look for paragraph breaks first (newline)
                paragraph_break = text.rfind('\n', search_start, end)
                if paragraph_break > search_start:
                    end = paragraph_break + 1
                else:
                    # Look for sentence endings
                    sentence_break = max(
                        text.rfind('. ', search_start, end),
                        text.rfind('.\n', search_start, end),
                        text.rfind('! ', search_start, end),
                        text.rfind('?\n', search_start, end)
                    )
                    if sentence_break > search_start:
                        end = sentence_break + 2
                    # Otherwise use the character limit
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            if end >= len(text):
                break
            start = max(start + 1, end - overlap)
        
        logger.info(f"Split text into {len(chunks)} chunks of average size {len(text) // len(chunks) if chunks else 0}")
        
        return chunks
    
    @staticmethod
    def merge_chunk_results(chunk_results: List[str], agent_name: str) -> str:
        """
        Merge results from multiple text chunks into a coherent final result
        
        Args:
            chunk_results: List of results from processing individual chunks
            agent_name: Name of the agent for context
            
        Returns:
            str: Merged result
        """
        if not chunk_results:
            return "No results to merge."
        
        if len(chunk_results) == 1:
            return chunk_results[0]
        
        # Create a structured summary format
        merged_result = f"# {agent_name} Analysis Results\n\n"
        merged_result += f"This document was processed in {len(chunk_results)} parts. Below are the consolidated results:\n\n"
        
        for i, result in enumerate(chunk_results, 1):
            merged_result += f"## Section {i} Results\n\n{result}\n\n"
        
        # Add a brief summary if we have multiple sections
        if len(chunk_results) > 2:
            merged_result += "## Overall Summary\n\n"
            merged_result += "The document has been analyzed in multiple sections. "
            merged_result += "Please review each section above for detailed findings. "
            merged_result += f"Total sections processed: {len(chunk_results)}\n"
        
        return merged_result