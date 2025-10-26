"""
Integration test for DOCX processing with real DOCX file
"""
import pytest
from fastapi import UploadFile
import os
from service.docx_service import DocxService
import tempfile
import io


class MockUploadFile:
    """Mock UploadFile for testing with real file content"""
    def __init__(self, filename, content):
        self.filename = filename
        self._content = content
    
    async def read(self):
        return self._content
    
    async def seek(self, position):
        pass


class TestDocxIntegration:
    """Integration tests for DOCX processing"""
    
    @pytest.mark.asyncio
    async def test_extract_text_from_real_docx(self):
        """Test text extraction from a dynamically created DOCX file"""
        from docx import Document
        import io
        
        # Create a test document in memory
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph with some content.')
        doc.add_paragraph('This is another paragraph to test text extraction.')
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Name'
        table.cell(0, 1).text = 'Age'
        table.cell(1, 0).text = 'John'
        table.cell(1, 1).text = '30'
        
        # Save to bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        file_content = docx_buffer.getvalue()
        
        # Create mock upload file
        mock_file = MockUploadFile("test_sample.docx", file_content)
        
        # Extract text
        result = await DocxService.extract_text_from_docx(mock_file)
        
        # Verify extracted content
        assert "Test Document" in result
        assert "This is a test paragraph" in result
        assert "This is another paragraph" in result
        assert "Table 1" in result
        assert "Name | Age" in result
        assert "John | 30" in result
        
        print(f"Extracted text length: {len(result)} characters")
        print("Extracted text preview:")
        print(result[:500])
    
    @pytest.mark.asyncio
    async def test_chunk_large_docx_content(self):
        """Test chunking of large DOCX content"""
        from docx import Document
        import io
        
        # Create a larger test document
        doc = Document()
        doc.add_heading('Large Test Document', 0)
        
        # Add multiple paragraphs to create larger content
        for i in range(10):
            doc.add_paragraph(f'This is paragraph {i+1} with some substantial content to test text chunking functionality. ')
        
        # Save to bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        file_content = docx_buffer.getvalue()
        
        mock_file = MockUploadFile("test_sample.docx", file_content)
        
        # Extract text
        text = await DocxService.extract_text_from_docx(mock_file)
        
        # Test chunking with small chunk size to force chunking
        chunks = DocxService.chunk_text(text, chunk_size=100)
        
        # Verify chunking works
        assert len(chunks) >= 1
        
        # Verify all chunks are within size limits (with some tolerance for boundary logic)
        for chunk in chunks:
            assert len(chunk) <= 150  # Allow some flexibility
        
        # Verify we can reconstruct most of the original text
        reconstructed = " ".join(chunks)
        # Should have most of the original content (allowing for some overlap/boundary handling)
        assert len(reconstructed) >= len(text) * 0.8
    
    def test_merge_results_integration(self):
        """Test merging multiple chunk results"""
        chunk_results = [
            "Analysis of section 1: Contains introduction and overview.",
            "Analysis of section 2: Detailed methodology and approach.",
            "Analysis of section 3: Results and conclusions.",
        ]
        
        merged = DocxService.merge_chunk_results(chunk_results, "test-docx-agent")
        
        # Verify structure
        assert "test-docx-agent Analysis Results" in merged
        assert "Section 1 Results" in merged
        assert "Section 2 Results" in merged
        assert "Section 3 Results" in merged
        assert "Overall Summary" in merged
        assert "Total sections processed: 3" in merged
        
        # Verify all original content is included
        for result in chunk_results:
            assert result in merged


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])